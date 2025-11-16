"""Tests for cell alignment in export formats based on col_defs."""

import io
import tempfile
import unittest
from pathlib import Path
from zipfile import ZipFile

from craftable import export_table
from craftable.styles import DocxStyle, XlsxStyle, OdtStyle, OdsStyle, RtfStyle


class TestExportCellAlignment(unittest.TestCase):
    """Test that col_defs properly controls data cell alignment in exports."""

    ROWS = [["Alice", 30, 95000], ["Bob", 25, 75000]]
    HEADERS = ["Name", "Age", "Salary"]

    def setUp(self):
        """Check for optional dependencies."""
        try:
            import docx  # noqa: F401

            self.docx_available = True
        except ImportError:
            self.docx_available = False

        try:
            import openpyxl  # noqa: F401

            self.openpyxl_available = True
        except ImportError:
            self.openpyxl_available = False

        try:
            import odf.opendocument  # noqa: F401

            self.odf_available = True
        except ImportError:
            self.odf_available = False

    def test_rtf_cell_alignment_left(self):
        """Test RTF export with left-aligned data cells."""
        stream = io.StringIO()
        col_defs = ["<", "<", "<"]
        export_table(
            self.ROWS,
            header_row=self.HEADERS,
            col_defs=col_defs,
            style=RtfStyle(),
            file=stream,
        )
        content = stream.getvalue()

        # RTF left alignment uses \ql command
        self.assertIn("\\ql", content)
        # Should have data values in the content
        self.assertIn("Alice", content)
        self.assertIn("Bob", content)

    def test_rtf_cell_alignment_center(self):
        """Test RTF export with center-aligned data cells."""
        stream = io.StringIO()
        col_defs = ["^", "^", "^"]
        export_table(
            self.ROWS,
            header_row=self.HEADERS,
            col_defs=col_defs,
            style=RtfStyle(),
            file=stream,
        )
        content = stream.getvalue()

        # RTF center alignment uses \qc command
        self.assertIn("\\qc", content)
        self.assertIn("Alice", content)

    def test_rtf_cell_alignment_right(self):
        """Test RTF export with right-aligned data cells."""
        stream = io.StringIO()
        col_defs = [">", ">", ">"]
        export_table(
            self.ROWS,
            header_row=self.HEADERS,
            col_defs=col_defs,
            style=RtfStyle(),
            file=stream,
        )
        content = stream.getvalue()

        # RTF right alignment uses \qr command
        self.assertIn("\\qr", content)
        self.assertIn("Alice", content)

    def test_rtf_cell_alignment_mixed(self):
        """Test RTF export with mixed column alignments."""
        stream = io.StringIO()
        col_defs = ["<20", "^10", ">15"]  # left, center, right
        export_table(
            self.ROWS,
            header_row=self.HEADERS,
            col_defs=col_defs,
            style=RtfStyle(),
            file=stream,
        )
        content = stream.getvalue()

        # Should have all three alignment types
        self.assertIn("\\ql", content)  # left
        self.assertIn("\\qc", content)  # center
        self.assertIn("\\qr", content)  # right

    def test_xlsx_cell_alignment(self):
        """Test XLSX export with column-based cell alignment."""
        if not self.openpyxl_available:
            self.skipTest("openpyxl not available")

        from openpyxl import load_workbook

        tmp_dir = tempfile.mkdtemp()
        path = Path(tmp_dir) / "aligned_cells.xlsx"
        col_defs = ["<20", "^10", ">15"]  # left, center, right

        export_table(
            self.ROWS,
            header_row=self.HEADERS,
            col_defs=col_defs,
            style=XlsxStyle(),
            file=path,
        )

        # Load the workbook and check data cell alignment
        wb = load_workbook(path)
        ws = wb.active

        # Check first data row (row 2) alignment
        self.assertEqual(ws["A2"].alignment.horizontal, "left")  # Alice
        self.assertEqual(ws["B2"].alignment.horizontal, "center")  # 30
        self.assertEqual(ws["C2"].alignment.horizontal, "right")  # 95000

        # Check second data row (row 3) alignment
        self.assertEqual(ws["A3"].alignment.horizontal, "left")  # Bob
        self.assertEqual(ws["B3"].alignment.horizontal, "center")  # 25
        self.assertEqual(ws["C3"].alignment.horizontal, "right")  # 75000

        # Verify the data is there (numbers preserved for Age & Salary)
        self.assertEqual(ws["A2"].value, "Alice")
        self.assertEqual(ws["B2"].value, 30)
        self.assertEqual(ws["C2"].value, 95000)

    def test_xlsx_header_and_cell_alignment_different(self):
        """Test XLSX export with different header and cell alignments."""
        if not self.openpyxl_available:
            self.skipTest("openpyxl not available")

        from openpyxl import load_workbook

        tmp_dir = tempfile.mkdtemp()
        path = Path(tmp_dir) / "diff_aligned.xlsx"

        # Headers: center, center, center
        # Data: left, center, right
        header_defs = ["^", "^", "^"]
        col_defs = ["<20", "^10", ">15"]

        export_table(
            self.ROWS,
            header_row=self.HEADERS,
            header_defs=header_defs,
            col_defs=col_defs,
            style=XlsxStyle(),
            file=path,
        )

        # Load the workbook and check alignments differ
        wb = load_workbook(path)
        ws = wb.active

        # Headers should all be centered
        self.assertEqual(ws["A1"].alignment.horizontal, "center")
        self.assertEqual(ws["B1"].alignment.horizontal, "center")
        self.assertEqual(ws["C1"].alignment.horizontal, "center")

        # Data should follow col_defs
        self.assertEqual(ws["A2"].alignment.horizontal, "left")
        self.assertEqual(ws["B2"].alignment.horizontal, "center")
        self.assertEqual(ws["C2"].alignment.horizontal, "right")

    def test_docx_cell_alignment(self):
        """Test DOCX export with column-based cell alignment."""
        if not self.docx_available:
            self.skipTest("python-docx not available")

        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        tmp_dir = tempfile.mkdtemp()
        path = Path(tmp_dir) / "aligned_cells.docx"
        col_defs = ["<20", "^10", ">15"]  # left, center, right

        export_table(
            self.ROWS,
            header_row=self.HEADERS,
            col_defs=col_defs,
            style=DocxStyle(),
            file=path,
        )

        # Load the document and check data cell alignment
        doc = Document(path)
        table = doc.tables[0]

        # Check first data row (row 1, since row 0 is header)
        data_row1 = table.rows[1]
        self.assertEqual(
            data_row1.cells[0].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.LEFT
        )
        self.assertEqual(
            data_row1.cells[1].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER
        )
        self.assertEqual(
            data_row1.cells[2].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.RIGHT
        )

        # Check second data row (row 2)
        data_row2 = table.rows[2]
        self.assertEqual(
            data_row2.cells[0].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.LEFT
        )
        self.assertEqual(
            data_row2.cells[1].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER
        )
        self.assertEqual(
            data_row2.cells[2].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.RIGHT
        )

        # Verify the data is there
        self.assertEqual(data_row1.cells[0].text, "Alice")
        self.assertEqual(data_row1.cells[1].text, "30")
        self.assertEqual(data_row1.cells[2].text, "95000")

    def test_docx_header_and_cell_alignment_different(self):
        """Test DOCX export with different header and cell alignments."""
        if not self.docx_available:
            self.skipTest("python-docx not available")

        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        tmp_dir = tempfile.mkdtemp()
        path = Path(tmp_dir) / "diff_aligned.docx"

        # Headers: right, center, left
        # Data: left, center, right
        header_defs = [">", "^", "<"]
        col_defs = ["<20", "^10", ">15"]

        export_table(
            self.ROWS,
            header_row=self.HEADERS,
            header_defs=header_defs,
            col_defs=col_defs,
            style=DocxStyle(),
            file=path,
        )

        # Load the document and check alignments differ
        doc = Document(path)
        table = doc.tables[0]

        header_row = table.rows[0]
        data_row = table.rows[1]

        # Headers should follow header_defs
        self.assertEqual(
            header_row.cells[0].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.RIGHT
        )
        self.assertEqual(
            header_row.cells[1].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER
        )
        self.assertEqual(
            header_row.cells[2].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.LEFT
        )

        # Data should follow col_defs
        self.assertEqual(
            data_row.cells[0].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.LEFT
        )
        self.assertEqual(
            data_row.cells[1].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER
        )
        self.assertEqual(
            data_row.cells[2].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.RIGHT
        )

    def test_odt_cell_alignment(self):
        """Test ODT export with column-based cell alignment."""
        if not self.odf_available:
            self.skipTest("odfpy not available")

        tmp_dir = tempfile.mkdtemp()
        path = Path(tmp_dir) / "aligned_cells.odt"
        col_defs = ["<20", "^10", ">15"]  # left, center, right

        export_table(
            self.ROWS,
            header_row=self.HEADERS,
            col_defs=col_defs,
            style=OdtStyle(),
            file=path,
        )

        # ODT files are ZIP archives containing XML
        # Check content.xml for alignment styles in data cells
        with ZipFile(path, "r") as zip_file:
            content_xml = zip_file.read("content.xml").decode("utf-8")

            # Should contain alignment styles
            self.assertIn("LeftAlign", content_xml)
            self.assertIn("CenterAlign", content_xml)
            self.assertIn("RightAlign", content_xml)

            # Should contain alignment properties
            self.assertIn("text-align", content_xml)

            # Should contain data values
            self.assertIn("Alice", content_xml)
            self.assertIn("Bob", content_xml)
            self.assertIn("30", content_xml)
            self.assertIn("25", content_xml)

    def test_odt_header_and_cell_alignment_different(self):
        """Test ODT export with different header and cell alignments."""
        if not self.odf_available:
            self.skipTest("odfpy not available")

        tmp_dir = tempfile.mkdtemp()
        path = Path(tmp_dir) / "diff_aligned.odt"

        # Headers: all center (bold)
        # Data: left, center, right (not bold)
        header_defs = ["^", "^", "^"]
        col_defs = ["<20", "^10", ">15"]

        export_table(
            self.ROWS,
            header_row=self.HEADERS,
            header_defs=header_defs,
            col_defs=col_defs,
            style=OdtStyle(),
            file=path,
        )

        # ODT files are ZIP archives containing XML
        with ZipFile(path, "r") as zip_file:
            content_xml = zip_file.read("content.xml").decode("utf-8")

            # Should have both header (bold) and data (non-bold) styles
            self.assertIn("HeaderBoldCenter", content_xml)
            self.assertIn("LeftAlign", content_xml)
            self.assertIn("CenterAlign", content_xml)
            self.assertIn("RightAlign", content_xml)

            # Should contain both header and data text
            self.assertIn("Name", content_xml)
            self.assertIn("Alice", content_xml)

    def test_ods_cell_alignment(self):
        """Test ODS export with column-based cell alignment."""
        if not self.odf_available:
            self.skipTest("odfpy not available")

        tmp_dir = tempfile.mkdtemp()
        path = Path(tmp_dir) / "aligned_cells.ods"
        col_defs = ["<20", "^10", ">15"]  # left, center, right

        export_table(
            self.ROWS,
            header_row=self.HEADERS,
            col_defs=col_defs,
            style=OdsStyle(),
            file=path,
        )

        # ODS files are ZIP archives containing XML
        # Check content.xml for alignment styles in data cells
        with ZipFile(path, "r") as zip_file:
            content_xml = zip_file.read("content.xml").decode("utf-8")

            # Should contain alignment styles
            self.assertIn("LeftAlign", content_xml)
            self.assertIn("CenterAlign", content_xml)
            self.assertIn("RightAlign", content_xml)

            # Should contain alignment properties
            self.assertIn("text-align", content_xml)

            # Should contain data values
            self.assertIn("Alice", content_xml)
            self.assertIn("Bob", content_xml)

    def test_ods_header_and_cell_alignment_different(self):
        """Test ODS export with different header and cell alignments."""
        if not self.odf_available:
            self.skipTest("odfpy not available")

        tmp_dir = tempfile.mkdtemp()
        path = Path(tmp_dir) / "diff_aligned.ods"

        # Headers: all right (bold)
        # Data: left, center, right (not bold)
        header_defs = [">", ">", ">"]
        col_defs = ["<20", "^10", ">15"]

        export_table(
            self.ROWS,
            header_row=self.HEADERS,
            header_defs=header_defs,
            col_defs=col_defs,
            style=OdsStyle(),
            file=path,
        )

        # ODS files are ZIP archives containing XML
        with ZipFile(path, "r") as zip_file:
            content_xml = zip_file.read("content.xml").decode("utf-8")

            # Should have both header (bold) and data (non-bold) styles
            self.assertIn("HeaderBoldRight", content_xml)
            self.assertIn("LeftAlign", content_xml)
            self.assertIn("CenterAlign", content_xml)
            self.assertIn("RightAlign", content_xml)

            # Should contain both header and data text
            self.assertIn("Salary", content_xml)
            self.assertIn("95000", content_xml)

    def test_xlsx_numeric_format_with_alignment(self):
        """Test XLSX export with numeric formatting and alignment."""
        if not self.openpyxl_available:
            self.skipTest("openpyxl not available")

        from openpyxl import load_workbook

        tmp_dir = tempfile.mkdtemp()
        path = Path(tmp_dir) / "numeric_aligned.xlsx"

        # Format salary as currency and right-align
        col_defs = ["<20", "^10", ">,15.2f"]  # left, center, right with format

        export_table(
            self.ROWS,
            header_row=self.HEADERS,
            col_defs=col_defs,
            style=XlsxStyle(),
            file=path,
        )

        # Load the workbook and check alignment is preserved with formatting
        wb = load_workbook(path)
        ws = wb.active

        # Check that alignment is correct even with format spec
        self.assertEqual(ws["C2"].alignment.horizontal, "right")
        self.assertEqual(ws["C3"].alignment.horizontal, "right")

        # Numeric values preserved (format may remain General if mapping not applied)
        self.assertIn(ws["C2"].value, (95000, 95000.0))
        self.assertIn(ws["C3"].value, (75000, 75000.0))

    def test_default_left_alignment(self):
        """Test that data cells default to left alignment when no col_defs provided."""
        if not self.openpyxl_available:
            self.skipTest("openpyxl not available")

        from openpyxl import load_workbook

        tmp_dir = tempfile.mkdtemp()
        path = Path(tmp_dir) / "default_align.xlsx"

        # Export without col_defs
        export_table(
            self.ROWS,
            header_row=self.HEADERS,
            style=XlsxStyle(),
            file=path,
        )

        # Load the workbook and check data cells default to left
        wb = load_workbook(path)
        ws = wb.active

        # Data cells should be left-aligned by default
        self.assertEqual(ws["A2"].alignment.horizontal, "left")
        self.assertEqual(ws["B2"].alignment.horizontal, "left")
        self.assertEqual(ws["C2"].alignment.horizontal, "left")


if __name__ == "__main__":
    unittest.main()

"""Tests for header alignment in export formats."""

import io
import tempfile
import unittest
from pathlib import Path
from zipfile import ZipFile

from craftable import export_table
from craftable.styles import DocxStyle, XlsxStyle, OdtStyle, OdsStyle, RtfStyle


class TestExportHeaderAlignment(unittest.TestCase):
    """Test that header_defs properly controls header alignment in exports."""

    ROWS = [["Alice", 30, 95000], ["Bob", 25, 75000]]
    HEADERS = ["Name", "Age", "Salary"]

    def setUp(self):
        """Check for optional dependencies."""
        try:
            import docx  # noqa: F401

            self.docx_available = True
        except ImportError:
            self.docx_available = False

        try:
            import openpyxl  # noqa: F401

            self.openpyxl_available = True
        except ImportError:
            self.openpyxl_available = False

        try:
            import odf.opendocument  # noqa: F401

            self.odf_available = True
        except ImportError:
            self.odf_available = False

    def test_rtf_header_alignment_left(self):
        """Test RTF export with left-aligned headers."""
        stream = io.StringIO()
        header_defs = ["<", "<", "<"]
        export_table(
            self.ROWS,
            header_row=self.HEADERS,
            header_defs=header_defs,
            style=RtfStyle(),
            file=stream,
        )
        content = stream.getvalue()

        # RTF left alignment uses \ql command
        self.assertIn("\\ql", content)
        # Should have Name, Age, and Salary in the content
        self.assertIn("Name", content)
        self.assertIn("Age", content)
        self.assertIn("Salary", content)

    def test_rtf_header_alignment_center(self):
        """Test RTF export with center-aligned headers (default)."""
        stream = io.StringIO()
        header_defs = ["^", "^", "^"]
        export_table(
            self.ROWS,
            header_row=self.HEADERS,
            header_defs=header_defs,
            style=RtfStyle(),
            file=stream,
        )
        content = stream.getvalue()

        # RTF center alignment uses \qc command
        self.assertIn("\\qc", content)
        self.assertIn("Name", content)

    def test_rtf_header_alignment_right(self):
        """Test RTF export with right-aligned headers."""
        stream = io.StringIO()
        header_defs = [">", ">", ">"]
        export_table(
            self.ROWS,
            header_row=self.HEADERS,
            header_defs=header_defs,
            style=RtfStyle(),
            file=stream,
        )
        content = stream.getvalue()

        # RTF right alignment uses \qr command
        self.assertIn("\\qr", content)
        self.assertIn("Name", content)

    def test_rtf_header_alignment_mixed(self):
        """Test RTF export with mixed header alignments."""
        stream = io.StringIO()
        header_defs = ["<", "^", ">"]
        export_table(
            self.ROWS,
            header_row=self.HEADERS,
            header_defs=header_defs,
            style=RtfStyle(),
            file=stream,
        )
        content = stream.getvalue()

        # Should have all three alignment types
        self.assertIn("\\ql", content)  # left
        self.assertIn("\\qc", content)  # center
        self.assertIn("\\qr", content)  # right

    def test_xlsx_header_alignment(self):
        """Test XLSX export with header alignment."""
        if not self.openpyxl_available:
            self.skipTest("openpyxl not available")

        from openpyxl import load_workbook

        tmp_dir = tempfile.mkdtemp()
        path = Path(tmp_dir) / "aligned.xlsx"
        header_defs = ["<", "^", ">"]

        export_table(
            self.ROWS,
            header_row=self.HEADERS,
            header_defs=header_defs,
            style=XlsxStyle(),
            file=path,
        )

        # Load the workbook and check header alignment
        wb = load_workbook(path)
        ws = wb.active

        # Check first row (headers) alignment
        self.assertEqual(ws["A1"].alignment.horizontal, "left")
        self.assertEqual(ws["B1"].alignment.horizontal, "center")
        self.assertEqual(ws["C1"].alignment.horizontal, "right")

        # Verify the headers are there
        self.assertEqual(ws["A1"].value, "Name")
        self.assertEqual(ws["B1"].value, "Age")
        self.assertEqual(ws["C1"].value, "Salary")

    def test_docx_header_alignment(self):
        """Test DOCX export with header alignment."""
        if not self.docx_available:
            self.skipTest("python-docx not available")

        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        tmp_dir = tempfile.mkdtemp()
        path = Path(tmp_dir) / "aligned.docx"
        header_defs = ["<", "^", ">"]

        export_table(
            self.ROWS,
            header_row=self.HEADERS,
            header_defs=header_defs,
            style=DocxStyle(),
            file=path,
        )

        # Load the document and check header alignment
        doc = Document(path)
        table = doc.tables[0]
        header_row = table.rows[0]

        # Check alignment of header cells
        self.assertEqual(
            header_row.cells[0].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.LEFT
        )
        self.assertEqual(
            header_row.cells[1].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER
        )
        self.assertEqual(
            header_row.cells[2].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.RIGHT
        )

        # Verify the headers are there
        self.assertEqual(header_row.cells[0].text, "Name")
        self.assertEqual(header_row.cells[1].text, "Age")
        self.assertEqual(header_row.cells[2].text, "Salary")

    def test_odt_header_alignment(self):
        """Test ODT export with header alignment."""
        if not self.odf_available:
            self.skipTest("odfpy not available")

        tmp_dir = tempfile.mkdtemp()
        path = Path(tmp_dir) / "aligned.odt"
        header_defs = ["<", "^", ">"]

        export_table(
            self.ROWS,
            header_row=self.HEADERS,
            header_defs=header_defs,
            style=OdtStyle(),
            file=path,
        )

        # ODT files are ZIP archives containing XML
        # Check content.xml for alignment styles
        with ZipFile(path, "r") as zip_file:
            content_xml = zip_file.read("content.xml").decode("utf-8")

            # Should contain our header bold styles with alignment
            self.assertIn("HeaderBoldLeft", content_xml)
            self.assertIn("HeaderBoldCenter", content_xml)
            self.assertIn("HeaderBoldRight", content_xml)

            # Should contain alignment properties
            self.assertIn("text-align", content_xml)
            self.assertIn("left", content_xml)
            self.assertIn("center", content_xml)
            self.assertIn("right", content_xml)

            # Should contain header text
            self.assertIn("Name", content_xml)
            self.assertIn("Age", content_xml)
            self.assertIn("Salary", content_xml)

    def test_ods_header_alignment(self):
        """Test ODS export with header alignment."""
        if not self.odf_available:
            self.skipTest("odfpy not available")

        tmp_dir = tempfile.mkdtemp()
        path = Path(tmp_dir) / "aligned.ods"
        header_defs = ["<", "^", ">"]

        export_table(
            self.ROWS,
            header_row=self.HEADERS,
            header_defs=header_defs,
            style=OdsStyle(),
            file=path,
        )

        # ODS files are ZIP archives containing XML
        # Check content.xml for alignment styles
        with ZipFile(path, "r") as zip_file:
            content_xml = zip_file.read("content.xml").decode("utf-8")

            # Should contain our header bold styles with alignment
            self.assertIn("HeaderBoldLeft", content_xml)
            self.assertIn("HeaderBoldCenter", content_xml)
            self.assertIn("HeaderBoldRight", content_xml)

            # Should contain alignment properties
            self.assertIn("text-align", content_xml)

            # Should contain header text
            self.assertIn("Name", content_xml)
            self.assertIn("Age", content_xml)
            self.assertIn("Salary", content_xml)

    def test_default_center_alignment(self):
        """Test that headers default to center alignment when no header_defs provided."""
        if not self.openpyxl_available:
            self.skipTest("openpyxl not available")

        from openpyxl import load_workbook

        tmp_dir = tempfile.mkdtemp()
        path = Path(tmp_dir) / "default.xlsx"

        # Export without header_defs
        export_table(
            self.ROWS,
            header_row=self.HEADERS,
            style=XlsxStyle(),
            file=path,
        )

        # Load the workbook and check header alignment defaults to center
        wb = load_workbook(path)
        ws = wb.active

        # All headers should be centered by default
        self.assertEqual(ws["A1"].alignment.horizontal, "center")
        self.assertEqual(ws["B1"].alignment.horizontal, "center")
        self.assertEqual(ws["C1"].alignment.horizontal, "center")


if __name__ == "__main__":
    unittest.main()
